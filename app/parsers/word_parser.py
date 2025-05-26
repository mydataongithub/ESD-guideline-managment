# app/parsers/word_parser.py
import io
from typing import Dict, Any, List, Optional
from docx import Document
import re
from PIL import Image
from .base_parser import BaseParser


class WordParser(BaseParser):
    """Parser to extract rule information from Word documents."""
    
    def __init__(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None):
        super().__init__(file_path, file_content)
        # Load document from file_path or file_content
        if file_path:
            self.doc = Document(file_path)
        else:
            self.doc = Document(io.BytesIO(file_content))
            
    def extract_rules(self) -> List[Dict[str, Any]]:
        """
        Extract rules from the Word document.
        Uses paragraph styles and formatting to identify rule-like content.
        
        Returns:
            List of dictionaries with rule data
        """
        rules = []
        paragraphs = self.doc.paragraphs
        
        # Try to identify rule structures based on headings and paragraph styles
        current_rule = None
        
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
                
            # Check if paragraph looks like a rule title
            if self._is_rule_title(para):
                # If we were building a previous rule, save it
                if current_rule and current_rule.get("content"):
                    rules.append(current_rule)
                
                # Start a new rule
                current_rule = {
                    "title": text,
                    "content": "",
                    "rule_type": self._detect_rule_type(text)
                }
            
            # If we have a current rule and this isn't a heading, add to content
            elif current_rule and not self._is_heading(para):
                if current_rule["content"]:
                    current_rule["content"] += "\n\n" + text
                else:
                    current_rule["content"] = text
        
        # Don't forget the last rule
        if current_rule and current_rule.get("content"):
            rules.append(current_rule)
            
        # If we didn't find any rules with the style-based approach, try based on text patterns
        if not rules:
            rules = self._extract_rules_by_patterns(paragraphs)
            
        return rules
    
    def _is_rule_title(self, paragraph) -> bool:
        """
        Check if paragraph looks like a rule title.
        
        Args:
            paragraph: docx Paragraph object
            
        Returns:
            True if likely a rule title
        """
        # Check style
        is_heading = paragraph.style.name.startswith('Heading')
        
        # Check formatting (bold text often indicates titles)
        is_bold = any(run.bold for run in paragraph.runs if run.bold is not None)
        
        # Check for rule number patterns
        text = paragraph.text.strip()
        has_rule_pattern = bool(re.match(r'(Rule\s+\d+|^\d+\.\d+|^\d+\.)\s+', text))
        
        # Check if reasonably short (titles usually aren't long)
        is_short = len(text) < 150
        
        return is_short and (is_heading or is_bold or has_rule_pattern)
    
    def _is_heading(self, paragraph) -> bool:
        """
        Check if paragraph is a heading.
        
        Args:
            paragraph: docx Paragraph object
            
        Returns:
            True if it's a heading
        """
        return paragraph.style.name.startswith('Heading')
    
    def _detect_rule_type(self, text: str) -> str:
        """
        Detect rule type based on text.
        
        Args:
            text: Text to analyze
            
        Returns:
            Rule type (esd, latchup, or general)
        """
        text_lower = text.lower()
        
        if "esd" in text_lower:
            return "esd"
        elif any(term in text_lower for term in ["latchup", "latch-up", "latch up"]):
            return "latchup"
        return "general"
    
    def _extract_rules_by_patterns(self, paragraphs) -> List[Dict[str, Any]]:
        """
        Extract rules by looking for text patterns.
        
        Args:
            paragraphs: List of document paragraphs
            
        Returns:
            List of rule dictionaries
        """
        rules = []
        full_text = "\n".join(p.text for p in paragraphs)
        
        # Look for rule patterns
        # Pattern: "Rule X: [Title]" followed by content
        rule_pattern = re.compile(r'Rule\s+(\d+):?\s+(.*?)(?:\n|$)(.*?)(?=Rule\s+\d+:|$)', re.DOTALL)
        for match in rule_pattern.finditer(full_text):
            rule_number, title, content = match.groups()
            content = content.strip()
            if title and content:
                rule = {
                    "title": f"Rule {rule_number}: {title.strip()}",
                    "content": content,
                    "rule_type": self._detect_rule_type(title + " " + content)
                }
                rules.append(rule)
        
        # Alternative pattern if no rules found
        if not rules:
            # Look for sections with numbering that might be rules
            section_pattern = re.compile(r'(\d+\.\d+\s+.*?)(?:\n|$)(.*?)(?=\d+\.\d+\s+|$)', re.DOTALL)
            for match in section_pattern.finditer(full_text):
                title, content = match.groups()
                content = content.strip()
                if title and content:
                    rule = {
                        "title": title.strip(),
                        "content": content,
                        "rule_type": self._detect_rule_type(title + " " + content)
                    }
                    rules.append(rule)
        
        return rules
    
    def extract_metadata(self) -> Dict[str, Any]:
        """
        Extract metadata from the Word document.
        
        Returns:
            Dictionary with metadata
        """
        metadata = {}
        props = self.doc.core_properties
        
        # Extract common properties
        if props.title:
            metadata["title"] = props.title
        if props.subject:
            metadata["subject"] = props.subject
        if props.author:
            metadata["author"] = props.author
        if props.created:
            metadata["created"] = props.created
        if props.modified:
            metadata["modified"] = props.modified
        if props.comments:
            metadata["comments"] = props.comments
            
        # Add document statistics
        metadata["paragraphs"] = len(self.doc.paragraphs)
        metadata["sections"] = len(self.doc.sections)
        
        return metadata
    
    def extract_images(self) -> List[Dict[str, Any]]:
        """
        Extract images from the Word document.
        
        Returns:
            List of dictionaries with image data
        """
        images = []
        
        # Word stores images as relationships in the document
        image_parts = []
        
        # Get image parts from the document
        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_parts.append(rel.target_part)
                except:
                    # Skip invalid relationships
                    pass
        
        # Process each image
        for i, img_part in enumerate(image_parts):
            try:
                # Get image content type and extension
                content_type = img_part.content_type
                ext = content_type.split("/")[-1]
                
                image_data = {
                    "filename": f"image_{i}.{ext}",
                    "image_data": img_part.blob,
                    "mime_type": content_type
                }
                images.append(image_data)
            except:
                # Skip images that can't be processed
                pass
        
        return images
